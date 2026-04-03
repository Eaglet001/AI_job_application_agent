from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from schemas import TailoredCV, CoverLetter
import os


def _add_section_heading(doc: Document, title: str):
    heading = doc.add_heading(title, level=2)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # Professional blue


def _add_divider(doc: Document):
    doc.add_paragraph("─" * 60)


def generate_tailored_cv_docx(cv: TailoredCV, output_path: str = "tailored_cv.docx") -> str:
    """Generates a formatted DOCX file from a TailoredCV object."""
    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Header: Name ─────────────────────────────────────────────────────────
    title = doc.add_heading(cv.name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)

    # ── Contact Info ─────────────────────────────────────────────────────────
    contact_parts = []
    if cv.email:
        contact_parts.append(cv.email)
    if cv.phone:
        contact_parts.append(cv.phone)
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)

    _add_divider(doc)

    # ── Professional Summary ─────────────────────────────────────────────────
    _add_section_heading(doc, "Professional Summary")
    doc.add_paragraph(cv.tailored_summary)

    # ── Skills ───────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Skills")
    skills_para = doc.add_paragraph()
    skills_para.add_run(" • ".join(cv.skills))

    # ── Experience ───────────────────────────────────────────────────────────
    _add_section_heading(doc, "Work Experience")
    for exp in cv.experience:
        role_para = doc.add_paragraph()
        role_para.add_run(f"{exp.role}").bold = True
        role_para.add_run(f" — {exp.company}  |  {exp.duration}")
        for resp in exp.responsibilities:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(resp)
            bullet.runs[0].font.size = Pt(10)

    # ── Education ────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Education")
    for edu in cv.education:
        edu_para = doc.add_paragraph()
        edu_para.add_run(f"{edu.degree}").bold = True
        edu_para.add_run(f" — {edu.institution}, {edu.year}")

    # ── Certifications ───────────────────────────────────────────────────────
    if cv.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in cv.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_cover_letter_docx(name: str, cover_letter: CoverLetter, output_path: str = "cover_letter.docx") -> str:
    """Generates a formatted DOCX cover letter."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    heading = doc.add_heading("Cover Letter", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = doc.add_paragraph(name)
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.runs[0].bold = True

    doc.add_paragraph("")

    for paragraph in cover_letter.content.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path
